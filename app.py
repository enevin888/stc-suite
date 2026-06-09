
**app.py**
```python
import difflib
import html
import io
import re

import streamlit as st

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None


st.set_page_config(
    page_title="SmartTech SOO Revision Review",
    page_icon="ST",
    layout="wide",
    initial_sidebar_state="collapsed",
)

NAVY = "#020C3D"
ORANGE = "#EC7626"
LIGHT_GRAY = "#E3E6EB"
WHITE = "#FFFFFF"
DARK_TEXT = "#111827"
MUTED_TEXT = "#5B6472"


def read_text_file(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    uploaded_file.seek(0)
    name = uploaded_file.name.lower()
    data = uploaded_file.read()

    if name.endswith(".docx"):
        if docx is None:
            st.error("python-docx is required for Word files.")
            return ""

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    if name.endswith(".pdf"):
        return extract_pdf_text(io.BytesIO(data))

    return data.decode("utf-8", errors="ignore")


def extract_pdf_text(file_like) -> str:
    if pdfplumber is None:
        st.error("pdfplumber is required for PDF files.")
        return ""

    text_parts = []
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                text_parts.append(text)

    return "\n\n".join(text_parts)


def get_diff_counts(original: str, updated: str) -> dict:
    diff = list(difflib.ndiff(original.splitlines(), updated.splitlines()))
    added = sum(1 for line in diff if line.startswith("+ "))
    removed = sum(1 for line in diff if line.startswith("- "))
    unchanged = sum(1 for line in diff if line.startswith("  "))

    return {
        "added": added,
        "removed": removed,
        "changed": added + removed,
        "unchanged": unchanged,
    }


def make_html_side_by_side_diff(
    original: str,
    updated: str,
    context_only: bool,
    left_title="Original SOO",
    right_title="Updated SOO",
) -> str:
    return difflib.HtmlDiff(tabsize=4, wrapcolumn=115).make_table(
        fromlines=original.splitlines(),
        tolines=updated.splitlines(),
        fromdesc=left_title,
        todesc=right_title,
        context=context_only,
        numlines=5,
    )


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def find_added_removed_lines(original: str, updated: str):
    diff = list(difflib.ndiff(original.splitlines(), updated.splitlines()))
    added = [line[2:].strip() for line in diff if line.startswith("+ ") and line[2:].strip()]
    removed = [line[2:].strip() for line in diff if line.startswith("- ") and line[2:].strip()]
    return added, removed


def find_alarm_names(text: str) -> set:
    alarms = set()
    patterns = [
        r"\b[A-Z]{1,8}_[A-Z0-9]{1,10}\b",
        r"\bAlarm\d+\.\(?\d+\)?\b",
        r"\b[A-Z]{2,8}\s*-\s*[^.\n]+",
    ]

    alarm_keywords = [
        "ALM",
        "HI",
        "LO",
        "FAIL",
        "FAULT",
        "FREEZE",
        "ZP",
        "MAT",
        "SAT",
        "DAT",
        "CSAT",
    ]

    for pattern in patterns:
        for match in re.findall(pattern, text):
            cleaned = match.strip(" .:-")
            if any(key in cleaned.upper() for key in alarm_keywords):
                alarms.add(cleaned)

    return alarms


def find_setpoints(text: str) -> set:
    setpoints = set()
    patterns = [
        r"[^.\n]*(?:setpoint|adj\.|adj|deg f|degrees f|%rh|h2o|in\. wc|btu)[^.\n]*",
        r"[^.\n]*\d+(?:\.\d+)?\s*(?:f|%rh|%|\" h2o|in\. wc|btu / lbm)[^.\n]*",
    ]

    for pattern in patterns:
        for match in re.findall(pattern, text, flags=re.IGNORECASE):
            cleaned = " ".join(match.strip().split())
            if len(cleaned) > 8:
                setpoints.add(cleaned)

    return setpoints


def detect_numbering_changes(added_lines, removed_lines):
    numbering_pattern = r"^\s*((?:#{1,6}\s*)?[A-Z]\.|\d+\.|[a-z]\.|\(\d+\)|\([a-z]\))"
    added_nums = [line for line in added_lines if re.search(numbering_pattern, line)]
    removed_nums = [line for line in removed_lines if re.search(numbering_pattern, line)]
    return added_nums, removed_nums


def make_engineering_comments(original: str, updated: str):
    comments = []
    added_lines, removed_lines = find_added_removed_lines(original, updated)

    original_alarms = find_alarm_names(original)
    updated_alarms = find_alarm_names(updated)
    added_alarms = sorted(updated_alarms - original_alarms)
    removed_alarms = sorted(original_alarms - updated_alarms)

    original_setpoints = find_setpoints(original)
    updated_setpoints = find_setpoints(updated)
    added_setpoints = sorted(updated_setpoints - original_setpoints)
    removed_setpoints = sorted(original_setpoints - updated_setpoints)

    added_numbering, removed_numbering = detect_numbering_changes(added_lines, removed_lines)

    def add(priority, category, comment, items=None):
        comments.append(
            {
                "priority": priority,
                "category": category,
                "comment": comment,
                "items": items or [],
            }
        )

    if added_alarms:
        add(
            "Review",
            "Added alarms",
            "New alarm references were added. Confirm trigger, delay, annunciation, and reset/clear logic.",
            added_alarms[:12],
        )

    if removed_alarms:
        add(
            "High",
            "Removed alarms",
            "Alarm references appear to have been removed. Confirm these were intentionally deleted.",
            removed_alarms[:12],
        )

    if added_setpoints or removed_setpoints:
        add(
            "Review",
            "Setpoint / threshold changes",
            "Setpoint or adjustable-value wording changed. Confirm values match design intent.",
            added_setpoints[:6] + removed_setpoints[:6],
        )

    if added_numbering or removed_numbering:
        add(
            "Review",
            "Numbering / formatting",
            "Numbered or lettered sequence lines changed. Verify numbering is clean and sequential.",
            added_numbering[:8] + removed_numbering[:8],
        )

    updated_lower = normalize_text(updated)

    if "shall modulate" in updated_lower and "set position" in updated_lower:
        add(
            "High",
            "Possible control contradiction",
            "Text includes both modulation and fixed-position language. Confirm the device is not described inconsistently.",
        )

    if "shared" in updated_lower and "output" in updated_lower:
        add(
            "Review",
            "Shared control output",
            "Shared output logic detected. Confirm single-unit, multi-unit, and disabled-unit behavior.",
        )

    if "damper" in updated_lower and "pressure" in updated_lower and "fan" in updated_lower:
        add(
            "Review",
            "Fan / damper coordination",
            "Fan, damper, and pressure control are referenced. Confirm only the intended device controls pressure.",
        )

    if "freeze" in updated_lower or "low temperature" in updated_lower:
        add(
            "Info",
            "Freeze protection",
            "Freeze / low-temperature protection exists. Confirm it overrides normal cooling/economizer control.",
        )

    if not comments:
        add(
            "Info",
            "No major flags detected",
            "No obvious alarm, setpoint, numbering, or contradiction flags were detected.",
        )

    return comments


def render_metric_row(metrics):
    html_out = '<div class="summary-row">'
    for icon, value, label, cls in metrics:
        html_out += f"""
        <div class="metric-card">
            <div class="metric-icon {cls}">{html.escape(str(icon))}</div>
            <div>
                <div class="metric-number">{html.escape(str(value))}</div>
                <div class="metric-label">{html.escape(str(label))}</div>
            </div>
        </div>
        """
    html_out += "</div>"
    st.markdown(html_out, unsafe_allow_html=True)


def render_comment_cards(comments):
    priority_colors = {"High": "#EF4444", "Review": ORANGE, "Info": "#64748B"}
    cards = ""

    for comment in comments:
        color = priority_colors.get(comment["priority"], "#64748B")
        items_html = ""

        if comment["items"]:
            safe_items = "".join(f"<li>{html.escape(str(item))}</li>" for item in comment["items"])
            items_html = f"<ul>{safe_items}</ul>"

        cards += f"""
        <div class="comment-card">
            <div class="comment-top">
                <span class="priority-pill" style="background:{color};">{html.escape(comment["priority"])}</span>
                <span class="comment-category">{html.escape(comment["category"])}</span>
            </div>
            <div class="comment-text">{html.escape(comment["comment"])}</div>
            {items_html}
        </div>
        """

    return cards


def render_comments_panel(title, comments):
    st.markdown(
        f"""
        <div class="review-panel">
            <div class="panel-title">{html.escape(title)}</div>
            <div class="comment-grid">
                {render_comment_cards(comments)}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_diff_panel(original_text, updated_text):
    context_only = st.toggle(
        "Show only changed areas",
        value=True,
        help="Turn off to show the full document.",
    )

    html_diff = make_html_side_by_side_diff(
        original_text,
        updated_text,
        context_only=context_only,
    )

    st.markdown(
        f"""
        <div class="diff-panel">
            <div class="panel-title">Line Changes</div>
            <div class="legend">
                <span class="legend-added">Added</span>
                <span class="legend-removed">Removed</span>
                <span class="legend-changed">Changed Within Line</span>
            </div>
            <div class="diff-wrapper">
                {html_diff}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def file_uploader_card(title, help_text, key, types):
    st.markdown(
        f"""
        <div class="upload-card">
            <div class="upload-title">{html.escape(title)}</div>
            <div class="upload-help">{html.escape(help_text)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    return st.file_uploader(
        title,
        type=types,
        key=key,
        label_visibility="collapsed",
    )


st.markdown(
    f"""
    <style>
        html, body, [class*="css"] {{ font-family: "Segoe UI", Arial, sans-serif; }}
        .stApp {{ background: {LIGHT_GRAY}; }}
        section[data-testid="stSidebar"] {{ display: none !important; }}
        .main .block-container {{
            padding-top: 1.25rem;
            padding-left: 2rem;
            padding-right: 2rem;
            max-width: 98%;
        }}

        .brand-header {{
            background: linear-gradient(90deg, {WHITE} 0%, #F8FAFC 72%, rgba(236,118,38,0.24) 100%);
            border: 1px solid #D7DCE5;
            border-radius: 18px;
            padding: 20px 26px;
            margin-bottom: 18px;
            box-shadow: 0 5px 18px rgba(2,12,61,0.10);
        }}

        .brand-row {{
            display: flex;
            align-items: center;
            gap: 24px;
        }}

        .fallback-logo {{
            font-size: 30px;
            font-weight: 950;
            color: {NAVY};
            padding-right: 24px;
            border-right: 2px solid #D7DCE5;
            line-height: 1.02;
            min-width: 190px;
        }}

        .fallback-logo span {{ color: {ORANGE}; }}

        .brand-title {{
            font-size: 38px;
            font-weight: 950;
            color: {NAVY};
            margin-bottom: 5px;
            line-height: 1.05;
        }}

        .brand-subtitle {{
            font-size: 16px;
            color: {MUTED_TEXT};
            font-weight: 650;
        }}

        .upload-card, .review-panel, .diff-panel {{
            background: {WHITE};
            border: 1px solid #D7DCE5;
            border-radius: 16px;
            padding: 18px;
            box-shadow: 0 5px 18px rgba(2,12,61,0.10);
            margin-top: 12px;
        }}

        .upload-card {{ border-top: 6px solid {ORANGE}; }}
        .upload-title {{ font-size: 22px; font-weight: 950; color: {NAVY}; margin-bottom: 4px; }}
        .upload-help {{ font-size: 14px; color: {MUTED_TEXT}; font-weight: 650; }}
        .panel-title {{ font-size: 25px; font-weight: 950; color: {NAVY}; margin-bottom: 10px; }}

        [data-testid="stFileUploader"] label {{ display: none; }}
        [data-testid="stFileUploaderDropzone"] {{
            background-color: #FAFBFC;
            border: 2px dashed #AAB4C3;
            border-radius: 14px;
            min-height: 92px;
            padding: 16px !important;
        }}

        [data-testid="stFileUploaderDropzone"]:hover {{
            border-color: {ORANGE};
            background-color: #FFF7F1;
        }}

        [data-testid="stFileUploaderDropzone"] * {{
            color: {NAVY} !important;
            font-weight: 850;
        }}

        .summary-row {{
            display: grid;
            grid-template-columns: repeat(4, minmax(160px, 1fr));
            gap: 14px;
            margin: 16px 0 18px 0;
        }}

        .metric-card {{
            background: {WHITE};
            border: 1px solid #D7DCE5;
            border-radius: 14px;
            padding: 16px 18px;
            box-shadow: 0 4px 14px rgba(2,12,61,0.08);
            display: flex;
            align-items: center;
            gap: 14px;
        }}

        .metric-icon {{
            width: 42px;
            height: 42px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-size: 24px;
            font-weight: 950;
        }}

        .metric-icon.added {{ background: #22A95A; }}
        .metric-icon.removed {{ background: #EF4444; }}
        .metric-icon.changed {{ background: {ORANGE}; }}
        .metric-icon.flags {{ background: {NAVY}; }}
        .metric-number {{ font-size: 30px; font-weight: 950; color: {NAVY}; line-height: 1; }}
        .metric-label {{ font-size: 13px; color: {MUTED_TEXT}; margin-top: 4px; font-weight: 700; }}

        .comment-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(310px, 1fr));
            gap: 12px;
        }}

        .comment-card {{
            border: 1px solid #D7DCE5;
            border-radius: 14px;
            padding: 14px 15px;
            background: #FAFBFC;
        }}

        .comment-top {{
            display: flex;
            align-items: center;
            gap: 9px;
            margin-bottom: 8px;
        }}

        .priority-pill {{
            color: white;
            font-size: 12px;
            font-weight: 950;
            padding: 4px 9px;
            border-radius: 999px;
        }}

        .comment-category {{ color: {NAVY}; font-size: 15px; font-weight: 950; }}
        .comment-text {{ color: {DARK_TEXT}; font-size: 14px; font-weight: 600; line-height: 1.35; }}
        .comment-card ul {{ margin-bottom: 0; padding-left: 18px; font-size: 12px; color: {MUTED_TEXT}; font-weight: 600; }}

        .legend {{ display: flex; gap: 8px; flex-wrap: wrap; margin-bottom: 12px; }}
        .legend span {{
            padding: 6px 12px;
            border-radius: 999px;
            font-size: 13px;
            font-weight: 900;
            color: {DARK_TEXT};
            border: 1px solid rgba(0,0,0,0.12);
        }}

        .legend-added {{ background-color: #B7F7C1; }}
        .legend-removed {{ background-color: #FFB6BD; }}
        .legend-changed {{ background-color: #FFE88A; }}

        .diff-wrapper {{
            background: {WHITE};
            border: 1px solid #CCD3DE;
            border-radius: 14px;
            overflow-x: auto;
            max-height: 74vh;
            overflow-y: auto;
        }}

        table.diff {{
            font-family: Consolas, "Courier New", monospace !important;
            border-collapse: collapse;
            width: 100%;
            font-size: 14px;
            color: {DARK_TEXT} !important;
            background: {WHITE} !important;
        }}

        .diff_header {{
            background-color: {NAVY} !important;
            color: white !important;
            font-weight: 950 !important;
            text-align: left;
            padding: 10px;
            position: sticky;
            top: 0;
            z-index: 2;
        }}

        td.diff_header {{
            width: 50px;
            text-align: right;
            color: {NAVY} !important;
            background-color: #E6EAF0 !important;
            font-weight: 800;
            border-right: 1px solid #B5BDCA;
            position: static;
        }}

        .diff_next {{ background-color: #F3F5F8 !important; color: {NAVY} !important; font-weight: 900; text-align: center; }}
        .diff_add {{ background-color: #B7F7C1 !important; color: #04210A !important; font-weight: 800; }}
        .diff_sub {{ background-color: #FFB6BD !important; color: #3B060B !important; font-weight: 800; }}
        .diff_chg {{ background-color: #FFE88A !important; color: #332600 !important; font-weight: 900; }}

        table.diff td {{
            padding: 7px 9px;
            vertical-align: top;
            border-bottom: 1px solid #E8ECF2;
            color: {DARK_TEXT} !important;
            white-space: pre-wrap;
            line-height: 1.38;
        }}

        div[data-testid="stToggle"] label {{
            color: {NAVY} !important;
            font-weight: 800 !important;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)


st.markdown(
    """
    <div class="brand-header">
        <div class="brand-row">
            <div class="fallback-logo"><span>smart</span> tech<br>contracting</div>
            <div>
                <div class="brand-title">SOO Revision Review</div>
                <div class="brand-subtitle">Compare original and updated Sequence of Operations documents.</div>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

col1, col2 = st.columns(2, gap="large")

with col1:
    original_file = file_uploader_card(
        "Original SOO",
        "Upload original SOO markdown, text, Word, or PDF.",
        "soo_original",
        ["md", "txt", "docx", "pdf"],
    )

with col2:
    updated_file = file_uploader_card(
        "Updated SOO",
        "Upload updated SOO markdown, text, Word, or PDF.",
        "soo_updated",
        ["md", "txt", "docx", "pdf"],
    )

if original_file and updated_file:
    original_text = read_text_file(original_file)
    updated_text = read_text_file(updated_file)

    if not original_text.strip() or not updated_text.strip():
        st.warning(
            "One or both files did not produce readable text. For PDFs, confirm the file has selectable text."
        )
        st.stop()

    counts = get_diff_counts(original_text, updated_text)
    comments = make_engineering_comments(original_text, updated_text)
    flag_count = sum(1 for c in comments if c["priority"] in ["High", "Review"])

    render_metric_row(
        [
            ("+", counts["added"], "Added Lines", "added"),
            ("-", counts["removed"], "Removed Lines", "removed"),
            ("~", counts["changed"], "Total Changed Lines", "changed"),
            ("!", flag_count, "Review Flags", "flags"),
        ]
    )

    render_comments_panel("Engineering Comments", comments)
    render_diff_panel(original_text, updated_text)
else:
    st.info("Upload the original SOO and updated SOO to begin.")
